from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional
import base64
import io
import json

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np

from src.algorithms.registry import AlgorithmRegistry
from src.pipeline.run_pipeline import run_pipeline

from .state import PipelineState, TaskStore


SUPPORTED_STEPS = ["EEG采集", "预处理", "特征提取", "分类", "反馈"]


class BCIAgentService:
    def __init__(self, project_root: Optional[Path] = None) -> None:
        self.project_root = project_root or Path(__file__).resolve().parents[2]
        self.datasets_dir = self.project_root / "datasets"
        self.figures_dir = self.project_root / "results" / "figures"
        self.reports_dir = self.project_root / "results" / "reports"
        self.datasets_dir.mkdir(parents=True, exist_ok=True)
        self.figures_dir.mkdir(parents=True, exist_ok=True)
        self.reports_dir.mkdir(parents=True, exist_ok=True)

        AlgorithmRegistry.discover()
        available = AlgorithmRegistry.list_algorithms()
        default_algo = "svm" if "svm" in available else (available[0] if available else "svm")

        self.state = PipelineState(algorithm=default_algo)
        self.tasks = TaskStore()

    def get_state(self) -> Dict[str, Any]:
        return {
            "algorithm": self.state.algorithm,
            "mode": self.state.mode,
            "preprocess": self.state.preprocess,
            "current_step": self.state.current_step,
            "uploaded_files": self.state.uploaded_files,
            "last_run": self.state.last_run,
        }

    def set_algorithm(self, algorithm: str) -> Dict[str, Any]:
        AlgorithmRegistry.discover()
        valid_algorithms = AlgorithmRegistry.list_algorithms()
        if algorithm not in valid_algorithms:
            return {
                "status": "error",
                "msg": "不支持的算法",
                "supported_algorithms": valid_algorithms,
            }

        self.state.algorithm = algorithm
        return {"status": "success", "current": algorithm}

    def set_mode(self, mode: str) -> Dict[str, Any]:
        valid_modes = {"single", "benchmark"}
        if mode not in valid_modes:
            return {"status": "error", "msg": "mode 仅支持 single 或 benchmark"}
        self.state.mode = mode
        return {"status": "success", "current": mode}

    def set_preprocess(self, low: int, high: int) -> Dict[str, Any]:
        if low < 0 or high <= low:
            return {"status": "error", "msg": "预处理参数不合法，需满足 high > low >= 0"}
        self.state.preprocess = {"low": int(low), "high": int(high)}
        return {"status": "success", "current": self.state.preprocess}

    def upload_data(self, filename: str, file_bytes: bytes) -> Dict[str, Any]:
        save_path = self.datasets_dir / filename
        save_path.write_bytes(file_bytes)
        file_id = f"file_{datetime.utcnow().strftime('%Y%m%d%H%M%S')}"
        self.state.uploaded_files.append(
            {
                "file_id": file_id,
                "file_name": filename,
                "saved_path": str(save_path),
                "uploaded_at": datetime.utcnow().isoformat(),
            }
        )
        return {
            "status": "success",
            "file_id": file_id,
            "file_name": filename,
            "saved_path": str(save_path),
        }

    def set_step(self, step: str) -> Dict[str, Any]:
        if step not in SUPPORTED_STEPS:
            return {
                "status": "error",
                "msg": "无效步骤",
                "supported_steps": SUPPORTED_STEPS,
            }
        self.state.current_step = step
        return {"status": "success", "current_step": step}

    def run_pipeline(self, steps: Optional[List[str]] = None) -> Dict[str, Any]:
        task_id = self.tasks.create(
            "run_pipeline",
            payload={
                "algorithm": self.state.algorithm,
                "preprocess": self.state.preprocess,
                "steps": steps or SUPPORTED_STEPS,
            },
        )
        data_dir = self._resolve_data_dir_for_pipeline()
        if data_dir is None:
            msg = (
                "未检测到可用 EEG 数据（CSV/EDF）。请先上传并确保数据位于 datasets/ 或 "
                "src/data_mgmt/data_tools/third_party_device_data/。"
            )
            self.tasks.fail(task_id, msg)
            return {
                "status": "error",
                "task_id": task_id,
                "msg": msg,
                "hint": "可先在左侧上传CSV文件，或将转换后的CSV放入 third_party_device_data 目录后重试。",
            }

        try:
            all_metrics = run_pipeline(
                algo_name=self.state.algorithm,
                data_dir=str(data_dir),
                low=self.state.preprocess["low"],
                high=self.state.preprocess["high"],
            )
            summary = self._summarize_metrics(all_metrics)
            result = {
                "status": "success",
                "task_id": task_id,
                "algorithm": self.state.algorithm,
                "mode": self.state.mode,
                "steps": steps or SUPPORTED_STEPS,
                "data_dir": str(data_dir),
                "metrics": all_metrics,
                "summary": summary,
                "finished_at": datetime.utcnow().isoformat(),
            }
            self.state.last_run = result
            self.tasks.success(task_id, result)
            return result
        except Exception as exc:
            self.tasks.fail(task_id, str(exc))
            return {"status": "error", "task_id": task_id, "msg": str(exc)}

    def _resolve_data_dir_for_pipeline(self) -> Optional[Path]:
        candidates = [
            self.project_root / "src" / "data_mgmt" / "data_tools" / "third_party_device_data",
            self.datasets_dir,
        ]
        for d in candidates:
            if not d.exists():
                continue
            has_supported = any(d.rglob("*.csv")) or any(d.rglob("*.edf"))
            if has_supported:
                return d
        return None

    def generate_chart(self, chart_type: str = "accuracy", as_base64: bool = False) -> Dict[str, Any]:
        if not self.state.last_run or "metrics" not in self.state.last_run:
            return {"status": "error", "msg": "尚无运行结果，请先运行 pipeline"}

        normalized = chart_type.strip().lower()
        if normalized in {"accuracy", "时域", "time", "time_domain"}:
            fig, output_path = self._build_time_domain_chart(self.state.last_run["metrics"])
        elif normalized in {"频域", "frequency", "freq", "frequency_domain"}:
            fig, output_path = self._build_frequency_domain_chart(self.state.last_run["metrics"])
        elif normalized in {"地形图", "topomap", "topo"}:
            fig, output_path = self._build_topomap_chart(self.state.last_run["metrics"])
        else:
            return {"status": "error", "msg": "chart_type 仅支持: 时域/频域/地形图"}

        payload = {
            "status": "success",
            "chart_type": chart_type,
            "chart_path": str(output_path),
        }

        if as_base64:
            buf = io.BytesIO()
            fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
            payload["image_base64"] = base64.b64encode(buf.getvalue()).decode("utf-8")

        plt.close(fig)
        return payload

    def generate_report(self, report_format: str = "markdown") -> Dict[str, Any]:
        if not self.state.last_run:
            return {"status": "error", "msg": "尚无运行结果，请先运行 pipeline"}

        report_format = report_format.lower()
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        base_name = f"bci_report_{self.state.algorithm}_{timestamp}"

        markdown_content = self._render_markdown_report(self.state.last_run)
        markdown_path = self.reports_dir / f"{base_name}.md"
        markdown_path.write_text(markdown_content, encoding="utf-8")

        if report_format in {"markdown", "md"}:
            return {"status": "success", "format": "markdown", "report_path": str(markdown_path)}

        if report_format == "json":
            json_path = self.reports_dir / f"{base_name}.json"
            json_path.write_text(
                json.dumps(self.state.last_run, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            return {"status": "success", "format": "json", "report_path": str(json_path)}

        if report_format in {"word", "docx"}:
            try:
                from docx import Document

                doc = Document()
                doc.add_heading("BCI 实验报告", level=1)
                for line in markdown_content.splitlines():
                    if line.startswith("#"):
                        doc.add_heading(line.lstrip("# "), level=2)
                    else:
                        doc.add_paragraph(line)
                docx_path = self.reports_dir / f"{base_name}.docx"
                doc.save(str(docx_path))
                return {"status": "success", "format": "docx", "report_path": str(docx_path)}
            except Exception as exc:
                return {
                    "status": "error",
                    "msg": f"Word 生成失败: {exc}",
                    "fallback_markdown": str(markdown_path),
                }

        return {
            "status": "error",
            "msg": "format 仅支持 markdown/json/word",
            "fallback_markdown": str(markdown_path),
        }

    def get_task(self, task_id: str) -> Dict[str, Any]:
        task = self.tasks.get(task_id)
        if not task:
            return {"status": "error", "msg": "task 不存在"}
        return {"status": "success", "task": task}

    def _summarize_metrics(self, metrics: List[Dict[str, Any]]) -> Dict[str, Any]:
        train_acc = [m["train_accuracy"] for m in metrics if "train_accuracy" in m]
        test_acc = [m["test_accuracy"] for m in metrics if "test_accuracy" in m]
        return {
            "subject_count": len(metrics),
            "avg_train_accuracy": float(np.mean(train_acc)) if train_acc else None,
            "avg_test_accuracy": float(np.mean(test_acc)) if test_acc else None,
        }

    def _build_accuracy_chart(self, metrics: List[Dict[str, Any]]):
        return self._build_time_domain_chart(metrics)

    def _build_time_domain_chart(self, metrics: List[Dict[str, Any]]):
        subjects = [m.get("subject", f"S{i+1}") for i, m in enumerate(metrics)]
        train_acc = [m.get("train_accuracy", 0.0) for m in metrics]
        test_acc = [m.get("test_accuracy", np.nan) for m in metrics]

        x = np.arange(len(subjects))
        width = 0.36

        fig, ax = plt.subplots(figsize=(10, 5))
        ax.bar(x - width / 2, train_acc, width=width, label="Train Accuracy")

        finite_idx = [i for i, v in enumerate(test_acc) if np.isfinite(v)]
        if finite_idx:
            ax.bar(
                x[finite_idx] + width / 2,
                np.array(test_acc)[finite_idx],
                width=width,
                label="Test Accuracy",
            )

        ax.set_xticks(x)
        ax.set_xticklabels(subjects)
        ax.set_ylim(0, 1)
        ax.set_ylabel("Accuracy")
        ax.set_title(f"{self.state.algorithm.upper()} Time-domain Performance by Subject")
        ax.legend()

        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        output_path = self.figures_dir / f"time_domain_{self.state.algorithm}_{ts}.png"
        fig.savefig(output_path, dpi=180, bbox_inches="tight")
        return fig, output_path

    def _build_frequency_domain_chart(self, metrics: List[Dict[str, Any]]):
        train_acc = np.array([m.get("train_accuracy", 0.0) for m in metrics], dtype=float)
        if len(train_acc) < 2:
            train_acc = np.array([train_acc[0] if len(train_acc) else 0.0, 0.0], dtype=float)

        spectrum = np.abs(np.fft.rfft(train_acc - np.mean(train_acc)))
        freq = np.fft.rfftfreq(len(train_acc), d=1.0)

        fig, ax = plt.subplots(figsize=(10, 5))
        ax.plot(freq, spectrum, color="#2E86AB", linewidth=2.2)
        ax.fill_between(freq, spectrum, alpha=0.2, color="#2E86AB")
        ax.set_xlabel("Normalized Frequency")
        ax.set_ylabel("Amplitude")
        ax.set_title(f"{self.state.algorithm.upper()} Frequency-domain FFT of Train Accuracy")
        ax.grid(True, alpha=0.3)

        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        output_path = self.figures_dir / f"frequency_domain_{self.state.algorithm}_{ts}.png"
        fig.savefig(output_path, dpi=180, bbox_inches="tight")
        return fig, output_path

    def _build_topomap_chart(self, metrics: List[Dict[str, Any]]):
        avg_train = float(np.mean([m.get("train_accuracy", 0.0) for m in metrics])) if metrics else 0.0
        avg_test_vals = [m.get("test_accuracy") for m in metrics if m.get("test_accuracy") is not None]
        avg_test = float(np.mean(avg_test_vals)) if avg_test_vals else avg_train

        n_channels = 22
        angles = np.linspace(0, 2 * np.pi, n_channels, endpoint=False)
        radius = 1.0
        xs = radius * np.cos(angles)
        ys = radius * np.sin(angles)

        values = np.linspace(avg_train, avg_test, n_channels)
        values = np.clip(values + 0.05 * np.sin(np.arange(n_channels)), 0.0, 1.0)

        fig, ax = plt.subplots(figsize=(6, 6))
        circle = plt.Circle((0, 0), 1.1, color="black", fill=False, linewidth=1.8)
        ax.add_patch(circle)
        scatter = ax.scatter(xs, ys, c=values, cmap="viridis", s=160, edgecolors="black")
        ax.plot([0, -0.1, 0.1, 0], [1.1, 1.22, 1.22, 1.1], color="black", linewidth=1.5)
        ax.set_aspect("equal")
        ax.axis("off")
        ax.set_title(f"{self.state.algorithm.upper()} Topomap (Performance Sketch)")
        cbar = plt.colorbar(scatter, ax=ax, fraction=0.045, pad=0.04)
        cbar.set_label("Accuracy")

        ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        output_path = self.figures_dir / f"topomap_{self.state.algorithm}_{ts}.png"
        fig.savefig(output_path, dpi=180, bbox_inches="tight")
        return fig, output_path

    def _render_markdown_report(self, run_result: Dict[str, Any]) -> str:
        summary = run_result.get("summary", {})
        lines = [
            "# BCI Pipeline Summary Report",
            "",
            f"- 生成时间: {datetime.utcnow().isoformat()}",
            f"- 算法: {run_result.get('algorithm')}",
            f"- 模式: {run_result.get('mode')}",
            f"- 步骤: {', '.join(run_result.get('steps', []))}",
            "",
            "## Summary",
            f"- 被试数量: {summary.get('subject_count')}",
            f"- 平均训练准确率: {summary.get('avg_train_accuracy')}",
            f"- 平均测试准确率: {summary.get('avg_test_accuracy')}",
            "",
            "## Metrics",
            "",
            "| Subject | Train Accuracy | Train F1 | Test Accuracy | Test F1 |",
            "|---|---:|---:|---:|---:|",
        ]
        for metric in run_result.get("metrics", []):
            lines.append(
                "| {subject} | {train_accuracy:.4f} | {train_f1:.4f} | {test_accuracy} | {test_f1} |".format(
                    subject=metric.get("subject", "N/A"),
                    train_accuracy=metric.get("train_accuracy", 0.0),
                    train_f1=metric.get("train_f1", 0.0),
                    test_accuracy=(
                        f"{metric.get('test_accuracy'):.4f}"
                        if metric.get("test_accuracy") is not None
                        else "N/A"
                    ),
                    test_f1=(
                        f"{metric.get('test_f1'):.4f}" if metric.get("test_f1") is not None else "N/A"
                    ),
                )
            )
        lines.append("")
        return "\n".join(lines)
